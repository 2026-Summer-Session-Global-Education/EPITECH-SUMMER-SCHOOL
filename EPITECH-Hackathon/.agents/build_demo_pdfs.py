from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BUILD_DIR = ROOT / ".agents" / "sample_pdf_build"
BUILD_DIR.mkdir(parents=True, exist_ok=True)

FONT = "Malgun Gothic"
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 103, 115)


DOCUMENTS = [
    {
        "filename": "01_해커톤_개최_안내",
        "kicker": "행사 공고",
        "title": "2026 청소년 사이버안전 해커톤 개최 안내",
        "subtitle": "미래디지털안전연구원 공식 행사 공고",
        "meta": [
            ("주최", "미래디지털안전연구원"),
            ("행사명", "2026 청소년 사이버안전 해커톤"),
            ("일시", "2026년 8월 20일 09:00-18:00"),
            ("장소", "한빛창업센터 3층"),
        ],
        "sections": [
            (
                "행사 목적",
                [
                    "미래디지털안전연구원은 청소년이 실제 보안 문제를 협업으로 해결하도록 2026 청소년 사이버안전 해커톤을 개최한다.",
                    "참가자는 피싱 탐지, 안전한 인증, 개인정보 보호 가운데 하나를 선택해 작동 가능한 시제품을 제작한다.",
                ],
            ),
            (
                "참가 및 산출물",
                [
                    "고등학생과 대학생으로 구성된 2-4인 팀이 참가할 수 있으며, 참가 신청 마감일은 2026년 8월 5일이다.",
                    "각 팀은 문제 정의서, 시제품, 5분 발표 자료를 제출해야 한다. 우수 팀은 미래디지털안전연구원장상을 받는다.",
                ],
            ),
        ],
    },
    {
        "filename": "02_해커톤_참가_협조_요청",
        "kicker": "협조 공문",
        "title": "2026 청소년 사이버안전 해커톤 참가 협조 요청",
        "subtitle": "교육기관 대상 참가자 추천 요청",
        "meta": [
            ("발신", "미래디지털안전연구원"),
            ("수신", "전국 고등학교 및 대학교"),
            ("관련 행사", "2026 청소년 사이버안전 해커톤"),
            ("추천 기한", "2026년 8월 5일"),
        ],
        "sections": [
            (
                "요청 사항",
                [
                    "미래디지털안전연구원은 2026 청소년 사이버안전 해커톤의 참가팀을 모집하오니 소속 학생을 추천해 주시기 바란다.",
                    "추천 학생은 2-4인 팀을 구성하고 피싱 탐지, 안전한 인증 또는 개인정보 보호 과제를 선택해야 한다.",
                ],
            ),
            (
                "참가 절차",
                [
                    "학교 담당자는 참가 신청서와 개인정보 이용 동의서를 2026년 8월 5일까지 제출한다.",
                    "선발된 팀은 2026년 8월 20일 한빛창업센터 3층에서 열리는 본 행사에 참가한다.",
                ],
            ),
        ],
    },
    {
        "filename": "03_사전교육_운영안",
        "kicker": "교육 운영안",
        "title": "해커톤 사전 보안교육 운영안",
        "subtitle": "본 행사 참가를 위한 필수 준비 과정",
        "meta": [
            ("운영", "미래디지털안전연구원"),
            ("대상", "2026 청소년 사이버안전 해커톤 참가팀"),
            ("교육일", "2026년 8월 13일"),
            ("방식", "온라인 실습"),
        ],
        "sections": [
            (
                "교육 목표",
                [
                    "이 교육은 2026 청소년 사이버안전 해커톤 참가팀이 시제품 개발 전에 이수해야 하는 필수 사전 과정이다.",
                    "미래디지털안전연구원 강사는 위협 모델링, 안전한 인증 설계, 개인정보 최소 수집 원칙을 실습으로 지도한다.",
                ],
            ),
            (
                "본 행사와의 연결",
                [
                    "교육 중 작성한 위협 모델 문서는 본 행사 문제 정의서의 필수 입력 자료로 사용된다.",
                    "사전교육 이수 확인을 받은 팀만 2026년 8월 20일 최종 해커톤에 참가할 수 있다.",
                ],
            ),
        ],
    },
    {
        "filename": "04_멘토링_가이드",
        "kicker": "운영 가이드",
        "title": "해커톤 현장 멘토링 가이드",
        "subtitle": "참가팀 지원을 위한 동반 자료",
        "meta": [
            ("발행", "미래디지털안전연구원"),
            ("적용 행사", "2026 청소년 사이버안전 해커톤"),
            ("적용일", "2026년 8월 20일"),
            ("대상", "기술 멘토 및 운영진"),
        ],
        "sections": [
            (
                "멘토 역할",
                [
                    "본 가이드는 2026 청소년 사이버안전 해커톤 현장에서 참가팀을 지원하는 기술 멘토의 동반 자료이다.",
                    "멘토는 정답을 대신 제시하지 않고 팀이 사전교육에서 작성한 위협 모델을 시제품 설계에 적용하도록 질문한다.",
                ],
            ),
            (
                "검토 기준",
                [
                    "멘토는 문제 정의의 명확성, 개인정보 최소 수집, 인증 실패 처리, 시연 가능성을 점검한다.",
                    "중간 점검 결과는 미래디지털안전연구원 운영본부에 전달하며 최종 심사 점수에는 직접 반영하지 않는다.",
                ],
            ),
        ],
    },
    {
        "filename": "05_해커톤_결과보고서",
        "kicker": "결과 보고",
        "title": "2026 청소년 사이버안전 해커톤 결과보고서",
        "subtitle": "행사 종료 후 성과 및 후속 조치",
        "meta": [
            ("주관", "미래디지털안전연구원"),
            ("행사명", "2026 청소년 사이버안전 해커톤"),
            ("개최일", "2026년 8월 20일"),
            ("참가 규모", "18개 팀, 62명"),
        ],
        "sections": [
            (
                "주요 결과",
                [
                    "미래디지털안전연구원은 2026년 8월 20일 한빛창업센터에서 2026 청소년 사이버안전 해커톤을 예정대로 완료했다.",
                    "18개 팀이 피싱 탐지, 안전한 인증, 개인정보 보호 시제품을 발표했으며 모든 팀이 문제 정의서와 시연 영상을 제출했다.",
                ],
            ),
            (
                "후속 조치",
                [
                    "최우수 팀의 피싱 탐지 시제품은 2026년 9월부터 미래디지털안전연구원 후속 멘토링 프로그램으로 이어진다.",
                    "운영진은 사전 보안교육과 현장 멘토링이 완성도 향상에 기여했다고 평가했다.",
                ],
            ),
        ],
    },
    {
        "filename": "06_해양생태_조사계획",
        "kicker": "대조 문서",
        "title": "푸른만 연안 잘피숲 생태 조사계획",
        "subtitle": "해커톤 문서와 관계가 없어야 하는 독립 샘플",
        "meta": [
            ("수행기관", "푸른만해양생태센터"),
            ("조사명", "2026 잘피숲 계절 변화 조사"),
            ("조사기간", "2026년 4월-10월"),
            ("조사지역", "푸른만 동부 연안"),
        ],
        "sections": [
            (
                "조사 목적",
                [
                    "푸른만해양생태센터는 연안 잘피숲의 면적 변화와 어린 물고기 서식 밀도를 계절별로 측정한다.",
                    "조사팀은 수온, 염분, 수중 가시거리와 잘피 피복률을 동일 지점에서 반복 기록한다.",
                ],
            ),
            (
                "조사 방법",
                [
                    "잠수 조사원은 50미터 고정 구간을 촬영하고 표본 사진을 식생 분석 시스템에 등록한다.",
                    "최종 결과는 연안 복원 우선순위를 정하는 기초 자료로 사용되며 사이버보안 행사나 교육 프로그램과는 관련이 없다.",
                ],
            ),
        ],
    },
]


def set_cell_font(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_bottom_border(paragraph, color="2E74B5", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_text(doc, text, size=11, bold=False, color=None, after=6, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_cell_font(run, size=size, bold=bold, color=color)
    return paragraph


def build_document(spec):
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 14, 6),
        ("Heading 2", 13, BLUE, 10, 4),
        ("Heading 3", 12, DARK, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("RELATIONSHIP GRAPH DEMO  |  SYNTHETIC SAMPLE")
    set_cell_font(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("시연용으로 제작된 합성 문서 - 실제 기관 또는 행사와 무관함")
    set_cell_font(run, size=8, color=MUTED)

    add_text(doc, spec["kicker"].upper(), size=9, bold=True, color=BLUE, after=3)
    add_text(doc, spec["title"], size=23, bold=True, color=RGBColor(20, 33, 48), after=4)
    subtitle = add_text(doc, spec["subtitle"], size=11.5, color=MUTED, after=12)
    add_bottom_border(subtitle)

    for label, value in spec["meta"]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        label_run = paragraph.add_run(f"{label}: ")
        set_cell_font(label_run, size=10.5, bold=True, color=DARK)
        value_run = paragraph.add_run(value)
        set_cell_font(value_run, size=10.5)

    for heading, paragraphs in spec["sections"]:
        doc.add_heading(heading, level=1)
        for text in paragraphs:
            add_text(doc, text, size=11, after=7)

    note = add_text(
        doc,
        "문서 식별 메모: 이 파일의 제목, 기관명, 행사명과 문장 단위 근거를 관계 그래프 시연에 사용할 수 있습니다.",
        size=9,
        color=MUTED,
        after=0,
    )
    note.paragraph_format.space_before = Pt(10)

    output_path = BUILD_DIR / f"{spec['filename']}.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    paths = [build_document(spec) for spec in DOCUMENTS]
    for path in paths:
        print(path)
